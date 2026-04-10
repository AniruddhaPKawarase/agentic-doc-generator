"""
services/document_generator.py  —  Word (.docx) document generation.

Converts the LLM markdown output into a professional Word document:
  - Project header with logo placeholder
  - Formatted tables (if present in the LLM output)
  - Section headings mapped from markdown ## / ###
  - Bullet points
  - Footer with generation metadata

Dependencies: python-docx
"""

import asyncio
import os
import re
import logging
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import get_settings
from models.schemas import GeneratedDocument

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentGenerator:
    """Generates a .docx file from LLM markdown output."""

    def __init__(self):
        self._docs_dir = Path(settings.docs_dir)
        self._docs_dir.mkdir(parents=True, exist_ok=True)

    async def generate(
        self,
        content: str,
        project_id: int,
        trade: str,
        document_type: str,
        project_name: str = "",
        title: str = None,
        set_ids: list = None,
        set_names: list = None,
        source_index: dict = None,
    ) -> GeneratedDocument:
        """
        Parse markdown content and produce a Word document (async entry point).

        Delegates all synchronous python-docx work to generate_sync() which is
        called via asyncio.to_thread().  This keeps the event loop unblocked
        during the ~100–500 ms of CPU/file-I/O work so that concurrent requests
        are not queued behind a single document build.
        """
        return await asyncio.to_thread(
            self.generate_sync,
            content=content,
            project_id=project_id,
            trade=trade,
            document_type=document_type,
            project_name=project_name,
            title=title,
            set_ids=set_ids,
            set_names=set_names,
            source_index=source_index,
        )

    def generate_sync(
        self,
        content: str,
        project_id: int,
        trade: str,
        document_type: str,
        project_name: str = "",
        title: str = None,
        set_ids: list = None,
        set_names: list = None,
        source_index: dict = None,
    ) -> GeneratedDocument:
        """
        Synchronous document builder — safe to call from a thread pool.
        Contains all python-docx operations (CPU + file I/O).
        """
        # Resolve display name: use provided project_name, or fall back to bare ID
        display_name = project_name or f"Project ID: {project_id}"

        doc = Document()
        self._configure_document(doc)
        self._add_header(
            doc, display_name, trade, document_type, title,
            set_ids=set_ids, set_names=set_names,
        )
        self._parse_and_add_content(doc, content)
        self._add_footer(doc, display_name, trade, set_names=set_names)

        if source_index and settings.source_ref_enabled:
            self._add_traceability_table(doc, source_index)

        # Build filename with project name slug
        # e.g. scope_electrical_GranvilleHotel_7298_a1b2c3d4.docx
        # With set_ids: scope_electrical_set4730_GranvilleHotel_7298_a1b2c3d4.docx
        file_id = str(uuid.uuid4())
        safe_trade = re.sub(r"[^\w\-]", "_", trade)
        safe_type = re.sub(r"[^\w\-]", "_", document_type)
        name_slug = self._project_name_slug(display_name, project_id)

        set_slug = ""
        if set_ids:
            set_slug = "set" + "_".join(str(s) for s in set_ids) + "_"
        filename = f"{safe_type}_{safe_trade}_{set_slug}{name_slug}_{project_id}_{file_id[:8]}.docx"
        # --- S3 MODE: save to temp, upload to S3, delete local ---
        if settings.storage_backend == "s3":
            import tempfile
            import sys as _sys
            _sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
            from s3_utils.operations import upload_file
            from s3_utils.helpers import generated_document_key

            # Save to temp file (python-docx needs a file path)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            doc.save(str(tmp_path))
            size_bytes = tmp_path.stat().st_size

            s3_key = generated_document_key(
                settings.s3_agent_prefix,
                project_name,
                project_id,
                trade,
                filename,
            )
            upload_ok = upload_file(str(tmp_path), s3_key)
            if upload_ok:
                logger.info(
                    "S3 upload OK: s3://%s/%s (%d bytes) | file_id=%s",
                    settings.s3_bucket_name, s3_key, size_bytes, file_id,
                )
            else:
                logger.error(
                    "S3 upload FAILED: s3://%s/%s | file_id=%s | "
                    "Check AWS credentials, bucket permissions, and network connectivity. "
                    "Verify STORAGE_BACKEND, S3_BUCKET_NAME, AWS_ACCESS_KEY_ID are in os.environ.",
                    settings.s3_bucket_name, s3_key, file_id,
                )

            # Clean up temp file — S3 is the only copy now
            tmp_path.unlink(missing_ok=True)
            download_url = f"{settings.docs_base_url}/{file_id}/download"
            file_path_str = f"s3://{settings.s3_bucket_name}/{s3_key}"

        # --- LOCAL MODE: save to docs_dir (original behavior) ---
        else:
            file_path = self._docs_dir / filename
            doc.save(str(file_path))
            size_bytes = file_path.stat().st_size
            download_url = f"{settings.docs_base_url}/{file_id}/download"
            file_path_str = str(file_path)

        meta = GeneratedDocument(
            file_id=file_id,
            filename=filename,
            file_path=file_path_str,
            download_url=download_url,
            project_id=project_id,
            trade=trade,
            document_type=document_type,
            size_bytes=size_bytes,
        )
        logger.info("Generated document: %s (%d bytes)", filename, size_bytes)
        return meta

    @staticmethod
    def _project_name_slug(display_name: str, project_id: int) -> str:
        """
        Extract a filename-safe slug from the project display name.

        "Granville Hotel (ID: 7298)"  →  "GranvilleHotel"
        "Project ID: 7298 (...)"      →  "project_7298"   (fallback)
        """
        # Strip the "(ID: XXXX ...)" suffix added by sql_service
        name = re.sub(r"\s*\(ID:\s*\d+.*$", "", display_name).strip()
        # If it's the bare "Project ID: XXXX" fallback, use a simple slug
        if name.upper().startswith("PROJECT ID:") or not name:
            return f"project_{project_id}"
        # Remove all non-alphanumeric characters (spaces, hyphens, etc.)
        slug = re.sub(r"[^\w]", "", name.replace(" ", ""))
        return slug[:30] if slug else f"project_{project_id}"

    def get_file_path(self, file_id: str) -> Path | None:
        """Find a generated file by its UUID."""
        for f in self._docs_dir.glob(f"*_{file_id[:8]}.docx"):
            return f
        return None

    # ── Hyperlinks and traceability ───────────────────────────────

    def _add_hyperlink(self, paragraph, url: str, text: str, color: str = "0563C1"):
        """Add a clickable hyperlink to a Word paragraph."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._element.append(hyperlink)

    def _add_traceability_table(self, doc, source_index):
        """Append a source reference traceability table to the document."""
        if not source_index:
            return
        doc.add_page_break()
        doc.add_heading("Source Reference Table", level=2)
        table = doc.add_table(rows=1, cols=4)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            table.style = "Table Grid"
        headers = ["Drawing Name", "Drawing Title", "PDF Link", "Coordinates"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for ref in sorted(source_index.values(), key=lambda r: r.drawing_name):
            row = table.add_row().cells
            p = row[0].paragraphs[0]
            if ref.s3_url:
                self._add_hyperlink(p, ref.s3_url, ref.drawing_name)
            else:
                p.text = ref.drawing_name
            row[1].text = ref.drawing_title
            p2 = row[2].paragraphs[0]
            if ref.s3_url:
                self._add_hyperlink(p2, ref.s3_url, "View PDF")
            else:
                p2.text = "N/A"
            if ref.x is not None and ref.y is not None:
                row[3].text = f"({ref.x}, {ref.y}) {ref.width}x{ref.height}"
            else:
                row[3].text = "\u2014"
        hyperlink_count = sum(1 for ref in source_index.values() if ref.s3_url)
        logger.info(
            "Traceability table: %d drawings, %d hyperlinks",
            len(source_index), hyperlink_count,
        )

    # ── Document structure builders ───────────────────────────────

    def _configure_document(self, doc: Document) -> None:
        """Set page margins and default font."""
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

        # Default font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    def _add_header(
        self, doc: Document, display_name: str, trade: str, document_type: str,
        title: str = None, set_ids: list = None, set_names: list = None,
    ) -> None:
        """Add a styled document header block."""
        # Company name / system name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("iFieldSmart — Construction Intelligence")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Document title
        doc_title = title or f"{document_type.title()} — {trade}"
        h = doc.add_heading(doc_title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        # Metadata row — shows "Project: Granville Hotel (ID: 7298)" everywhere
        meta_parts = [
            f"Project: {display_name}",
            f"Trade: {trade}",
        ]
        if set_names:
            meta_parts.append(f"Set: {', '.join(set_names)}")
        elif set_ids:
            meta_parts.append(f"Set ID(s): {', '.join(str(s) for s in set_ids)}")
        meta_parts.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_run = meta_para.add_run("  •  ".join(meta_parts))
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Divider
        doc.add_paragraph("─" * 80)

    def _parse_and_add_content(self, doc: Document, content: str) -> None:
        """
        Parse markdown-ish content and add to the document.
        Handles: ## headings, ### subheadings, bullets (- / *), tables, paragraphs.
        """
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Detect markdown table start (line with | separators)
            if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[-| :]+\|\s*$", lines[i + 1]):
                table_lines = [line]
                i += 1  # skip separator row
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table_from_markdown(doc, table_lines)
                continue

            # H2 heading
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)

            # H3 heading
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)

            # H4 heading
            elif line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=3)

            # Bullet
            elif re.match(r"^\s*[-*•]\s", line):
                text = re.sub(r"^\s*[-*•]\s+", "", line).strip()
                text = self._strip_inline_markdown(text)
                doc.add_paragraph(text, style="List Bullet")

            # Numbered list
            elif re.match(r"^\s*\d+\.\s", line):
                text = re.sub(r"^\s*\d+\.\s+", "", line).strip()
                text = self._strip_inline_markdown(text)
                doc.add_paragraph(text, style="List Number")

            # Empty line → paragraph break
            elif line.strip() == "":
                pass  # Skip blanks

            # Regular paragraph
            else:
                cleaned = self._strip_inline_markdown(line.strip())
                if cleaned:
                    doc.add_paragraph(cleaned)

            i += 1

    def _add_table_from_markdown(self, doc: Document, table_lines: list[str]) -> None:
        """Convert markdown table rows into a Word table."""
        rows = [
            [cell.strip() for cell in line.strip("|").split("|")]
            for line in table_lines
            if line.strip()
        ]
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data[:num_cols]):
                cell = row.cells[c_idx]
                cell.text = self._strip_inline_markdown(cell_text)
                if r_idx == 0:
                    # Bold header row
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # Blue header background
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "1E3A5F")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)

        doc.add_paragraph()  # spacing after table

    def _add_footer(self, doc: Document, display_name: str, trade: str, set_names: list = None) -> None:
        """Add document footer with metadata."""
        section = doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_parts = [
            "iFieldSmart Construction Intelligence",
            display_name,
            f"Trade: {trade}",
        ]
        if set_names:
            footer_parts.append(f"Set: {', '.join(set_names)}")
        footer_parts.append("AI-generated — verify against source drawings")

        run = para.add_run("  •  ".join(footer_parts))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    @staticmethod
    def _strip_inline_markdown(text: str) -> str:
        """Remove **bold** and *italic* markers but keep the text."""
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        return text
