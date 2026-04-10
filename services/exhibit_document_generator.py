"""
services/exhibit_document_generator.py — Generates professional exhibit-style
Word documents matching the format of the sample Exhibit Electrical Scope of Work.

Output structure (mirrors the sample .docx files):
  ┌──────────────────────────────────────┐
  │  PROJECT HEADER BLOCK                │
  │  (Project name, trade, date, etc.)   │
  ├──────────────────────────────────────┤
  │  EXHIBIT TITLE                       │
  ├──────────────────────────────────────┤
  │  SCOPE SUMMARY                       │
  ├──────────────────────────────────────┤
  │  SCOPE OF WORK BY DRAWING            │
  │  Drawing No | Notes table per sheet  │
  ├──────────────────────────────────────┤
  │  MASTER NOTES TABLE                  │
  │  Drawing No | Trade | CSI | Note     │
  └──────────────────────────────────────┘

Dependencies: python-docx
"""

from __future__ import annotations

import asyncio
import logging
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from config import get_settings
from models.schemas import GeneratedDocument

logger = logging.getLogger(__name__)
settings = get_settings()

# Brand colours (matching sample files)
COLOUR_DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)   # #1E3A5F — headings
COLOUR_MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6 — section bars
COLOUR_LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)  # #F2F2F2 — alt row shading
COLOUR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOUR_DARK_TEXT = RGBColor(0x26, 0x26, 0x26)
COLOUR_SUBTEXT   = RGBColor(0x6B, 0x72, 0x80)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_bold_colour(run, colour: RGBColor, size_pt: float = 11) -> None:
    run.bold = True
    run.font.color.rgb = colour
    run.font.size = Pt(size_pt)


def _strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"^#{1,4}\s+", "", text)
    return text.strip()


class ExhibitDocumentGenerator:
    """
    Generates a Word exhibit document styled to match the sample files.

    Call generate_sync() from a thread (python-docx is synchronous).
    Call generate() from async code — it wraps generate_sync in to_thread().
    """

    def __init__(self):
        self._docs_dir = Path(settings.docs_dir)
        self._docs_dir.mkdir(parents=True, exist_ok=True)

    # ── Public async entry point ───────────────────────────────────

    async def generate(
        self,
        content: str,
        project_name: str,
        trade: str,
        document_type: str,
        drawing_summary: Optional[list[dict]] = None,
        title: Optional[str] = None,
        source_index: dict = None,
    ) -> GeneratedDocument:
        return await asyncio.to_thread(
            self.generate_sync,
            content=content,
            project_name=project_name,
            trade=trade,
            document_type=document_type,
            drawing_summary=drawing_summary,
            title=title,
            source_index=source_index,
        )

    # ── Synchronous builder ────────────────────────────────────────

    def generate_sync(
        self,
        content: str,
        project_name: str,
        trade: str,
        document_type: str,
        drawing_summary: Optional[list[dict]] = None,
        title: Optional[str] = None,
        # project_id used for filename and GeneratedDocument metadata
        project_id: int = 0,
        source_index: dict = None,
    ) -> GeneratedDocument:
        # Resolve display name
        if project_id and not project_name:
            project_name = f"Project ID: {project_id}"
        display_name = project_name or f"Project ID: {project_id}"

        doc = Document()
        self._configure_document(doc)
        self._add_cover_block(doc, display_name, trade, document_type, title)
        self._add_scope_content(doc, content, trade)

        if drawing_summary:
            self._add_drawing_table(doc, drawing_summary, trade)

        self._add_footer(doc, display_name, trade)

        if source_index and settings.source_ref_enabled:
            self._add_traceability_table(doc, source_index)

        # Build filename with project name slug + project_id
        # e.g. Exhibit_GranvilleHotel_Electrical_scope_7298_a1b2c3d4.docx
        file_id = str(uuid.uuid4())
        safe_trade = re.sub(r"[^\w\-]", "_", trade)
        safe_type = re.sub(r"[^\w\-]", "_", document_type)
        name_slug = self._project_name_slug(display_name, project_id)
        filename = f"Exhibit_{name_slug}_{safe_trade}_{safe_type}_{project_id}_{file_id[:8]}.docx"
        # --- S3 MODE: save to temp, upload to S3, delete local ---
        if settings.storage_backend == "s3":
            import tempfile
            import sys as _sys
            _sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
            from s3_utils.operations import upload_file
            from s3_utils.helpers import generated_document_key

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = Path(tmp.name)
            doc.save(str(tmp_path))
            size_bytes = tmp_path.stat().st_size

            s3_key = generated_document_key(
                settings.s3_agent_prefix,
                project_name,
                project_id,
                trade,
                filename,
            )
            upload_ok = upload_file(str(tmp_path), s3_key)
            if upload_ok:
                logger.info(
                    "S3 upload OK (exhibit): s3://%s/%s (%d bytes) | file_id=%s",
                    settings.s3_bucket_name, s3_key, size_bytes, file_id,
                )
            else:
                logger.error(
                    "S3 upload FAILED (exhibit): s3://%s/%s | file_id=%s | "
                    "Check AWS credentials, bucket permissions, and network connectivity.",
                    settings.s3_bucket_name, s3_key, file_id,
                )

            tmp_path.unlink(missing_ok=True)
            download_url = f"{settings.docs_base_url}/{file_id}/download"
            file_path_str = f"s3://{settings.s3_bucket_name}/{s3_key}"

        # --- LOCAL MODE: save to docs_dir (original behavior) ---
        else:
            file_path = self._docs_dir / filename
            doc.save(str(file_path))
            size_bytes = file_path.stat().st_size
            download_url = f"{settings.docs_base_url}/{file_id}/download"
            file_path_str = str(file_path)

        meta = GeneratedDocument(
            file_id=file_id,
            filename=filename,
            file_path=file_path_str,
            download_url=download_url,
            project_id=project_id,
            trade=trade,
            document_type=document_type,
            size_bytes=size_bytes,
        )
        logger.info("Exhibit document generated: %s (%d bytes)", filename, size_bytes)
        return meta

    def get_file_path(self, file_id: str) -> Optional[Path]:
        for f in self._docs_dir.glob(f"*_{file_id[:8]}.docx"):
            return f
        return None

    # ── Hyperlinks and traceability ────────────────────────────────

    def _add_hyperlink(self, paragraph, url: str, text: str, color: str = "0563C1"):
        """Add a clickable hyperlink to a Word paragraph."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._element.append(hyperlink)

    def _add_traceability_table(self, doc, source_index):
        """Append a source reference traceability table to the document."""
        if not source_index:
            return
        doc.add_page_break()
        doc.add_heading("Source Reference Table", level=2)
        table = doc.add_table(rows=1, cols=4)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            table.style = "Table Grid"
        headers = ["Drawing Name", "Drawing Title", "PDF Link", "Coordinates"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for ref in sorted(source_index.values(), key=lambda r: r.drawing_name):
            row = table.add_row().cells
            p = row[0].paragraphs[0]
            if ref.s3_url:
                self._add_hyperlink(p, ref.s3_url, ref.drawing_name)
            else:
                p.text = ref.drawing_name
            row[1].text = ref.drawing_title
            p2 = row[2].paragraphs[0]
            if ref.s3_url:
                self._add_hyperlink(p2, ref.s3_url, "View PDF")
            else:
                p2.text = "N/A"
            if ref.x is not None and ref.y is not None:
                size = f" {ref.width}x{ref.height}" if ref.width is not None and ref.height is not None else ""
                row[3].text = f"({ref.x}, {ref.y}){size}"
            else:
                row[3].text = "\u2014"
        hyperlink_count = sum(1 for ref in source_index.values() if ref.s3_url)
        logger.info(
            "Traceability table: %d drawings, %d hyperlinks",
            len(source_index), hyperlink_count,
        )

    # ── Document structure builders ────────────────────────────────

    def _configure_document(self, doc: Document) -> None:
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    def _add_cover_block(
        self,
        doc: Document,
        project_name: str,
        trade: str,
        document_type: str,
        title: Optional[str],
    ) -> None:
        """Top section: company name, project, exhibit title, date."""

        # Company banner
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner.add_run("iFieldSmart  |  Construction Intelligence Platform")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = COLOUR_SUBTEXT

        # Divider line
        doc.add_paragraph("─" * 100)

        # Exhibit title (large, dark blue)
        exhibit_title = title or f"Exhibit — {trade} Scope of Work"
        h = doc.add_heading(exhibit_title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if h.runs:
            h.runs[0].font.color.rgb = COLOUR_DARK_BLUE
            h.runs[0].font.size = Pt(18)

        # Metadata block (2-column feel using tabs)
        meta_lines = [
            ("Project:", project_name),
            ("Trade:", trade),
            ("Document Type:", document_type.title()),
            ("Prepared:", datetime.utcnow().strftime("%B %d, %Y")),
            ("Prepared by:", "AI Construction Intelligence"),
            ("Status:", "DRAFT — Verify against source drawings"),
        ]
        for label, value in meta_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            label_run = p.add_run(f"{label:<20}")
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = COLOUR_DARK_BLUE
            val_run = p.add_run(value)
            val_run.font.size = Pt(10)
            val_run.font.color.rgb = COLOUR_DARK_TEXT

        doc.add_paragraph("─" * 100)
        doc.add_paragraph()  # spacer

    def _add_scope_content(self, doc: Document, content: str, trade: str) -> None:
        """Parse LLM markdown output into structured Word content."""
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Markdown table
            if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[-| :]+\|\s*$", lines[i + 1]):
                table_lines = [line]
                i += 2  # skip separator
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_markdown_table(doc, table_lines)
                continue

            # H1 / H2 heading
            if line.startswith("## ") or line.startswith("# "):
                text = re.sub(r"^#+\s+", "", line).strip()
                h = doc.add_heading(text, level=1)
                if h.runs:
                    h.runs[0].font.color.rgb = COLOUR_MID_BLUE
            # H3 heading
            elif line.startswith("### "):
                text = line[4:].strip()
                h = doc.add_heading(_strip_markdown(text), level=2)
                if h.runs:
                    h.runs[0].font.color.rgb = COLOUR_DARK_BLUE
            # H4
            elif line.startswith("#### "):
                text = line[5:].strip()
                h = doc.add_heading(_strip_markdown(text), level=3)
            # Bold drawing reference line (e.g., **Drawing E-101 — Electrical**)
            elif re.match(r"^\*\*Drawing\s+", line):
                text = _strip_markdown(line)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.color.rgb = COLOUR_DARK_BLUE
                run.font.size = Pt(11)
            # Bullet
            elif re.match(r"^\s*[-*•]\s", line):
                text = re.sub(r"^\s*[-*•]\s+", "", line).strip()
                text = _strip_markdown(text)
                if text:
                    doc.add_paragraph(text, style="List Bullet")
            # Numbered list
            elif re.match(r"^\s*\d+\.\s", line):
                text = re.sub(r"^\s*\d+\.\s+", "", line).strip()
                text = _strip_markdown(text)
                if text:
                    doc.add_paragraph(text, style="List Number")
            # Empty line → spacer
            elif line.strip() == "":
                pass
            # Regular paragraph
            else:
                cleaned = _strip_markdown(line)
                if cleaned:
                    doc.add_paragraph(cleaned)

            i += 1

    def _add_drawing_table(self, doc: Document, drawing_summary: list[dict], trade: str) -> None:
        """
        Add a structured "Scope by Drawing Number" table.

        Each row: Drawing No | Source Trade | CSI Division | Scope Notes
        Header row is dark-blue with white text. Alternating row shading.
        """
        doc.add_heading("Scope of Work — Drawing Reference Table", level=1)

        if not drawing_summary:
            doc.add_paragraph("No drawing data available.")
            return

        # Flatten to table rows
        flat_rows: list[tuple[str, str, str, str]] = []
        for entry in drawing_summary:
            drawing_no = entry.get("drawing_no", "")
            source_trade = entry.get("source_trade", "")
            csi = ", ".join(entry.get("csi", []))
            notes = entry.get("notes", [])
            for note in notes:
                flat_rows.append((drawing_no, source_trade, csi, note.strip()))
            if not notes:
                flat_rows.append((drawing_no, source_trade, csi, ""))

        if not flat_rows:
            return

        headers = ["Drawing No", "Source Trade", "CSI Division", "Scope / Note"]
        table = doc.add_table(rows=1 + len(flat_rows), cols=4)
        table.style = "Table Grid"

        # Widths: 1.0 | 1.3 | 2.0 | 5.0 inches
        col_widths = [Inches(1.0), Inches(1.3), Inches(2.0), Inches(5.0)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

        # Header row
        header_row = table.rows[0]
        for c_idx, hdr in enumerate(headers):
            cell = header_row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(hdr)
            _set_bold_colour(run, COLOUR_WHITE, size_pt=10)
            _set_cell_bg(cell, "2E75B6")

        # Data rows
        for r_idx, (drawing_no, source_trade, csi, note) in enumerate(flat_rows):
            row = table.rows[r_idx + 1]
            row.cells[0].text = drawing_no
            row.cells[1].text = source_trade
            row.cells[2].text = csi
            row.cells[3].text = note

            # Alternating row shading
            if r_idx % 2 == 0:
                for cell in row.cells:
                    _set_cell_bg(cell, "F2F2F2")

            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()  # spacing after table

    def _add_markdown_table(self, doc: Document, table_lines: list[str]) -> None:
        """Convert markdown pipe-table to Word table."""
        rows = [
            [cell.strip() for cell in line.strip("|").split("|")]
            for line in table_lines
            if line.strip()
        ]
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data[:num_cols]):
                cell = row.cells[c_idx]
                cell.text = _strip_markdown(cell_text)
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = COLOUR_WHITE
                    _set_cell_bg(cell, "2E75B6")
                elif r_idx % 2 == 1:
                    _set_cell_bg(cell, "F2F2F2")

                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

    @staticmethod
    def _project_name_slug(display_name: str, project_id: int) -> str:
        """Extract a filename-safe slug from the project display name."""
        name = re.sub(r"\s*\(ID:\s*\d+.*$", "", display_name).strip()
        if name.upper().startswith("PROJECT ID:") or not name:
            return f"project_{project_id}"
        slug = re.sub(r"[^\w]", "", name.replace(" ", ""))
        return slug[:30] if slug else f"project_{project_id}"

    def _add_footer(self, doc: Document, project_name: str, trade: str) -> None:
        section = doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            f"iFieldSmart Construction Intelligence  |  {project_name}  |  "
            f"Trade: {trade}  |  AI-generated — verify against source drawings  |  "
            f"{datetime.utcnow().strftime('%Y-%m-%d')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = COLOUR_SUBTEXT
