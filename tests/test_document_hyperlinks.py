"""Tests for Word document hyperlinks and traceability tables."""
import os
import sys
from pathlib import Path

# Force local storage so tests write to disk and can be read back
os.environ["STORAGE_BACKEND"] = "local"

# Ensure project root is on the path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

# Clear lru_cache so Settings re-reads env vars with local storage
import config as _config
_config.get_settings.cache_clear()

import pytest
from docx import Document
from services.source_index import SourceReference


def _make_source_index() -> dict[str, SourceReference]:
    from services.source_index import Annotation
    return {
        "A102": SourceReference(
            drawing_id=318845, drawing_name="A102",
            drawing_title="ARCH SITE PLAN",
            s3_url="https://bucket.s3.amazonaws.com/path/pdf.pdf",
            pdf_name="pdfA102", x=100, y=200, width=50, height=30,
            text="", annotations=(),
        ),
        "E-101": SourceReference(
            drawing_id=12345, drawing_name="E-101",
            drawing_title="ELECTRICAL FLOOR PLAN",
            s3_url="https://bucket.s3.amazonaws.com/path2/pdf2.pdf",
            pdf_name="pdfE101", x=300, y=400, width=60, height=40,
            text="", annotations=(),
        ),
    }


class TestDocumentGeneratorHyperlinks:
    def test_traceability_table_added(self):
        from services.document_generator import DocumentGenerator
        gen = DocumentGenerator()
        source_index = _make_source_index()
        result = gen.generate_sync(
            content="## Scope\n- Item 1\n- Item 2",
            project_id=7292, trade="Civil", document_type="scope",
            project_name="Test Project (ID: 7292)",
            source_index=source_index,
        )
        doc = Document(result.file_path)
        tables = doc.tables
        assert len(tables) >= 1
        last_table = tables[-1]
        assert len(last_table.rows) == 3  # header + 2 data rows

    def test_traceability_table_skipped_when_empty(self):
        from services.document_generator import DocumentGenerator
        gen = DocumentGenerator()
        result = gen.generate_sync(
            content="## Scope\n- Item 1",
            project_id=7292, trade="Civil", document_type="scope",
            project_name="Test Project (ID: 7292)",
            source_index={},
        )
        doc = Document(result.file_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Source Reference Table" not in headings

    def test_document_without_source_index(self):
        from services.document_generator import DocumentGenerator
        gen = DocumentGenerator()
        result = gen.generate_sync(
            content="## Scope\n- Item 1",
            project_id=7292, trade="Civil", document_type="scope",
            project_name="Test Project (ID: 7292)",
            source_index=None,
        )
        assert result is not None
        assert result.filename.endswith(".docx")

    def test_exhibit_traceability_table(self):
        from services.exhibit_document_generator import ExhibitDocumentGenerator
        gen = ExhibitDocumentGenerator()
        source_index = _make_source_index()
        result = gen.generate_sync(
            content="## Scope\n- Item 1",
            project_name="Test Project (ID: 7292)",
            trade="Civil", document_type="scope",
            project_id=7292,
            source_index=source_index,
        )
        doc = Document(result.file_path)
        tables = doc.tables
        assert len(tables) >= 1
