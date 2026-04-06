"""
scope_pipeline/services/export_service.py — Combined multi-trade Word document export.

Generates a single Word (.docx) report containing scope gap results for all
trades in a ProjectSession, optionally filtered to a subset of trades.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime, timezone
from typing import Optional

from scope_pipeline.models_v2 import ProjectSession

logger = logging.getLogger(__name__)

# Brand colour (RGB tuple matching existing document_agent.py conventions)
_DARK_BLUE = (0, 51, 102)
_MID_GRAY = (128, 128, 128)


class ExportService:
    """Generate combined multi-trade Word documents from a ProjectSession."""

    def __init__(self, docs_dir: str = "./generated_docs") -> None:
        self._docs_dir = docs_dir
        os.makedirs(self._docs_dir, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_combined_word(
        self,
        session: ProjectSession,
        trades: Optional[list[str]] = None,
    ) -> str:
        """Generate a single Word document containing results for all trades.

        Args:
            session: The project session holding per-trade pipeline results.
            trades:  Optional allowlist of trade names (case-insensitive).
                     When None all trades with results are included.

        Returns:
            Absolute path of the saved .docx file.
        """
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename = f"ScopeGap_{session.project_id}_AllTrades_{timestamp}.docx"
        path = os.path.join(self._docs_dir, filename)
        self._write_word(path, session, trades)
        logger.info("Combined Word document saved: %s", path)
        return path

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _resolve_trades(
        self,
        session: ProjectSession,
        trades: Optional[list[str]],
    ) -> list[str]:
        """Return the ordered list of trade keys to include in the document."""
        available = list(session.trade_results.keys())
        if trades is None:
            return available
        allowed = {t.lower() for t in trades}
        return [t for t in available if t.lower() in allowed]

    def _write_word(
        self,
        path: str,
        session: ProjectSession,
        trades: Optional[list[str]],
    ) -> None:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Default body font
        normal_style = doc.styles["Normal"]
        normal_style.font.size = Pt(10)
        normal_style.font.name = "Calibri"

        trade_keys = self._resolve_trades(session, trades)
        now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

        # ----------------------------------------------------------------
        # Cover / title page
        # ----------------------------------------------------------------
        title_para = doc.add_heading("SCOPE GAP REPORT \u2014 ALL TRADES", level=0)
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(*_DARK_BLUE)

        doc.add_paragraph(f"Project: {session.project_name} (ID: {session.project_id})")
        doc.add_paragraph(f"Generated: {now_str}")
        doc.add_paragraph(f"Trades included: {len(trade_keys)}")

        # ----------------------------------------------------------------
        # Per-trade sections
        # ----------------------------------------------------------------
        total_items = 0

        for trade_key in trade_keys:
            container = session.trade_results[trade_key]
            result = container.latest_result

            # Page break before every trade section
            _add_page_break(doc)

            trade_label = container.trade

            # Trade heading
            trade_heading = doc.add_heading(
                f"{trade_label} \u2014 Scope of Work", level=1
            )
            for run in trade_heading.runs:
                run.font.color.rgb = RGBColor(*_DARK_BLUE)

            if result is None:
                doc.add_paragraph("No results available for this trade.")
                continue

            # Summary counts
            item_count = len(result.items)
            ambiguity_count = len(result.ambiguities)
            gotcha_count = len(result.gotchas)
            completeness_pct = result.completeness.overall_pct
            total_items += item_count

            doc.add_paragraph(
                f"Items: {item_count}   |   "
                f"Ambiguities: {ambiguity_count}   |   "
                f"Gotchas: {gotcha_count}   |   "
                f"Completeness: {completeness_pct:.1f}%"
            )

            # Items grouped by drawing_name
            if result.items:
                doc.add_heading("Scope Items", level=2)
                grouped: dict[str, list] = {}
                for item in result.items:
                    grouped.setdefault(item.drawing_name, []).append(item)

                for drawing_name, drawing_items in grouped.items():
                    doc.add_heading(drawing_name, level=3)
                    for item in drawing_items:
                        p = doc.add_paragraph()
                        bold_run = p.add_run(item.text)
                        bold_run.bold = True
                        p.add_run(
                            f"\n  CSI: {item.csi_code}"
                            f" | Confidence: {item.confidence:.0%}"
                            f" | Source: \"{item.source_snippet}\""
                        )

            # Ambiguities
            if result.ambiguities:
                doc.add_heading("Ambiguities", level=2)
                for amb in result.ambiguities:
                    p = doc.add_paragraph()
                    p.add_run(amb.scope_text).bold = True
                    p.add_run(
                        f"\n  Competing trades: {', '.join(amb.competing_trades)}"
                        f" | Severity: {amb.severity}"
                        f"\n  Recommendation: {amb.recommendation}"
                    )

            # Gotchas
            if result.gotchas:
                doc.add_heading("Gotchas", level=2)
                for gtc in result.gotchas:
                    p = doc.add_paragraph()
                    p.add_run(f"[{gtc.severity.upper()}] {gtc.description}").bold = True
                    p.add_run(
                        f"\n  Risk type: {gtc.risk_type}"
                        f" | Affected trades: {', '.join(gtc.affected_trades)}"
                        f"\n  Recommendation: {gtc.recommendation}"
                    )

        # ----------------------------------------------------------------
        # Summary page
        # ----------------------------------------------------------------
        _add_page_break(doc)

        summary_heading = doc.add_heading("Report Summary", level=1)
        for run in summary_heading.runs:
            run.font.color.rgb = RGBColor(*_DARK_BLUE)

        doc.add_paragraph(f"Total trades: {len(trade_keys)}")
        doc.add_paragraph(f"Total scope items: {total_items}")
        doc.add_paragraph(f"Project: {session.project_name} (ID: {session.project_id})")
        doc.add_paragraph(f"Generated: {now_str}")

        # ----------------------------------------------------------------
        # Footer paragraph
        # ----------------------------------------------------------------
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            "Generated by iFieldSmart ScopeAI Pipeline v2.0"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(*_MID_GRAY)

        doc.save(path)


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def _add_page_break(doc) -> None:  # type: ignore[no-untyped-def]
    """Insert an explicit page break paragraph into *doc*."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
