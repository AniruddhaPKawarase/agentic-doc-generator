"""
scope_pipeline/services/document_agent.py -- Agent 7: Document generation (NO LLM).

Input: Validated pipeline results (items, ambiguities, gotchas, completeness, quality, stats).
Output: DocumentSet with file paths to Word, PDF, CSV, JSON files.

Clean-export redesign: Word/PDF/CSV emit professional scope-of-work documents
suitable for handing to a subcontractor.  JSON remains a full data dump.

All 4 formats are generated in parallel via asyncio.to_thread().
"""

from __future__ import annotations

import asyncio
import csv
import json
import logging
import os
import re
from datetime import datetime, timezone
from typing import Any, Optional, Union
from uuid import uuid4

from scope_pipeline.models import (
    AmbiguityItem,
    ClassifiedItem,
    CompletenessReport,
    DocumentSet,
    GotchaItem,
    PipelineStats,
    QualityReport,
)

logger = logging.getLogger(__name__)

# Brand colours
_DARK_BLUE = (0, 51, 102)       # #003366
_GRAY = (102, 102, 102)         # #666666


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insert a clickable hyperlink into a python-docx Paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "003366")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _build_filename(
    project_name: str,
    project_id: int,
    trade: str,
    ext: str,
    file_id: str = "",
) -> str:
    """Build a professional filename with unique ID for download API.

    Example: ``7276_Singh_Residence_Concrete_Scope_of_Work_a1b2c3d4.docx``
    """
    slug = re.sub(r"[^\w\s-]", "", project_name).strip()
    slug = re.sub(r"\s+", "_", slug).title()
    trade_slug = re.sub(r"[^\w\s-]", "", trade).strip()
    trade_slug = re.sub(r"\s+", "_", trade_slug).title()
    id_suffix = f"_{file_id[:8]}" if file_id else ""
    return f"{project_id}_{slug}_{trade_slug}_Scope_of_Work{id_suffix}.{ext}"


def _group_items_by_drawing(
    items: list[ClassifiedItem],
) -> dict[str, list[ClassifiedItem]]:
    """Group items by drawing_name, sorted alphabetically."""
    grouped: dict[str, list[ClassifiedItem]] = {}
    for item in items:
        grouped.setdefault(item.drawing_name, []).append(item)
    return dict(sorted(grouped.items()))


# ------------------------------------------------------------------
# Document Agent
# ------------------------------------------------------------------

class DocumentAgent:
    """Generate Word, PDF, CSV, JSON output files from validated pipeline results."""

    def __init__(self, docs_dir: str = "./generated_docs") -> None:
        self._docs_dir = docs_dir
        os.makedirs(self._docs_dir, exist_ok=True)
        try:
            from config import get_settings
            self._settings = get_settings()
        except Exception:
            self._settings = None

    async def generate_all(
        self,
        items: list[ClassifiedItem],
        ambiguities: list[AmbiguityItem],
        gotchas: list[GotchaItem],
        completeness: CompletenessReport,
        quality: QualityReport,
        project_id: int,
        project_name: str,
        trade: str,
        stats: PipelineStats,
        project_location: str = "",
        drawing_s3_urls: dict[str, str] | None = None,
        set_name: str = "",
        set_id: int = 0,
    ) -> DocumentSet:
        """Generate all 4 document formats in parallel. Returns DocumentSet.

        When STORAGE_BACKEND=s3, uploads generated files to S3 and returns
        download URLs compatible with ``GET /api/documents/{file_id}/download``.
        When STORAGE_BACKEND=local (or settings unavailable), returns local paths.
        """
        common_kwargs: dict[str, Any] = dict(
            items=items,
            ambiguities=ambiguities,
            gotchas=gotchas,
            completeness=completeness,
            quality=quality,
            project_id=project_id,
            project_name=project_name,
            trade=trade,
            stats=stats,
            project_location=project_location,
            drawing_s3_urls=drawing_s3_urls or {},
        )

        # Generate a unique file_id per format for download API resolution
        exts = ["docx", "pdf", "csv", "json"]
        file_ids = {ext: uuid4().hex for ext in exts}

        word_path = os.path.join(
            self._docs_dir, _build_filename(project_name, project_id, trade, "docx", file_ids["docx"]),
        )
        pdf_path = os.path.join(
            self._docs_dir, _build_filename(project_name, project_id, trade, "pdf", file_ids["pdf"]),
        )
        csv_path = os.path.join(
            self._docs_dir, _build_filename(project_name, project_id, trade, "csv", file_ids["csv"]),
        )
        json_path = os.path.join(
            self._docs_dir, _build_filename(project_name, project_id, trade, "json", file_ids["json"]),
        )

        word_task = asyncio.to_thread(self.generate_word_sync, word_path, **common_kwargs)
        pdf_task = asyncio.to_thread(self.generate_pdf_sync, pdf_path, **common_kwargs)
        csv_task = asyncio.to_thread(self.generate_csv_sync, csv_path, **common_kwargs)
        json_task = asyncio.to_thread(self.generate_json_sync, json_path, **common_kwargs)

        results = await asyncio.gather(
            word_task, pdf_task, csv_task, json_task, return_exceptions=True,
        )

        # Collect successfully generated files: attr -> (local_path, file_id)
        paths = [word_path, pdf_path, csv_path, json_path]
        attrs = ["word_path", "pdf_path", "csv_path", "json_path"]
        generated: dict[str, tuple[str, str]] = {}
        for result, path, attr, ext in zip(results, paths, attrs, exts):
            if isinstance(result, Exception):
                logger.error("Document generation failed for %s: %s", attr, result)
            else:
                generated[attr] = (path, file_ids[ext])

        # S3 mode: upload files and return download URLs
        use_s3 = (
            self._settings is not None
            and getattr(self._settings, "storage_backend", "local") == "s3"
        )
        if use_s3:
            return await self._upload_to_s3(
                generated, project_name, project_id, set_name, set_id, trade,
            )

        # Local mode: return local file paths (backward compatible)
        doc_set = DocumentSet()
        for attr, (path, _fid) in generated.items():
            setattr(doc_set, attr, path)
        return doc_set

    # ------------------------------------------------------------------
    # S3 upload
    # ------------------------------------------------------------------

    async def _upload_to_s3(
        self,
        generated: dict[str, tuple[str, str]],
        project_name: str,
        project_id: int,
        set_name: str,
        set_id: int,
        trade: str,
    ) -> DocumentSet:
        """Upload generated files to S3 and return DocumentSet with download URLs."""
        from s3_utils.operations import upload_file, delete_prefix
        from s3_utils.helpers import generated_document_key

        settings = self._settings
        docs_base_url = getattr(settings, "docs_base_url", "")
        agent_prefix = getattr(settings, "s3_agent_prefix", "construction-intelligence-agent")

        # Build the S3 folder prefix and clear old docs for this project/set/trade
        sample_key = generated_document_key(
            agent_prefix, project_name, project_id,
            set_name, set_id, trade, "placeholder",
        )
        folder_prefix = "/".join(sample_key.split("/")[:-1]) + "/"
        try:
            deleted = await asyncio.to_thread(delete_prefix, folder_prefix)
            if deleted > 0:
                logger.info(
                    "Overwrite: deleted %d old scope docs from %s", deleted, folder_prefix,
                )
        except Exception as exc:
            logger.warning("Failed to delete old scope docs at %s: %s", folder_prefix, exc)

        # Upload each file in parallel
        doc_set = DocumentSet()

        async def _upload_one(attr: str, local_path: str, file_id: str) -> None:
            filename = os.path.basename(local_path)
            s3_key = generated_document_key(
                agent_prefix, project_name, project_id,
                set_name, set_id, trade, filename,
            )
            ok = await asyncio.to_thread(upload_file, local_path, s3_key)
            if ok:
                download_url = f"{docs_base_url}/{file_id[:8]}/download"
                setattr(doc_set, attr, download_url)
                logger.info("S3 upload OK: %s → %s", filename, s3_key)
                # Clean up local temp file — S3 is the single source of truth
                try:
                    os.unlink(local_path)
                except OSError:
                    pass
            else:
                # Fallback: keep local path if S3 upload fails
                setattr(doc_set, attr, local_path)
                logger.error("S3 upload FAILED for %s — keeping local copy", filename)

        upload_tasks = [
            _upload_one(attr, local_path, file_id)
            for attr, (local_path, file_id) in generated.items()
        ]
        await asyncio.gather(*upload_tasks)

        return doc_set

    # ------------------------------------------------------------------
    # Word (.docx) -- clean scope-of-work export
    # ------------------------------------------------------------------

    def generate_word_sync(self, path: str, **kwargs: Any) -> str:
        items: list[ClassifiedItem] = kwargs["items"]
        project_id: int = kwargs["project_id"]
        project_name: str = kwargs["project_name"]
        trade: str = kwargs["trade"]
        project_location: str = kwargs.get("project_location", "")
        drawing_s3_urls: dict[str, str] = kwargs.get("drawing_s3_urls", {})

        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(10)
        style.font.name = "Calibri"

        # --- Title: centered, dark blue, 16pt bold ---
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f"SCOPE OF WORK \u2014 {trade.upper()}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(*_DARK_BLUE)

        # --- Project info lines: gray 11pt ---
        info_lines = [f"Project: {project_name} (ID: {project_id})"]
        if project_location:
            info_lines.append(f"Location: {project_location}")
        info_lines.append(f"Date: {datetime.now(timezone.utc).strftime('%B %d, %Y')}")

        for line in info_lines:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(*_GRAY)

        # --- Scope items grouped by drawing ---
        grouped = _group_items_by_drawing(items)

        for drawing_name, drawing_items in grouped.items():
            # Drawing heading -- hyperlink if S3 URL available
            s3_url = drawing_s3_urls.get(drawing_name)
            if s3_url:
                heading_para = doc.add_paragraph()
                heading_para.style = doc.styles["Heading 2"]
                _add_hyperlink(heading_para, s3_url, drawing_name)
            else:
                doc.add_heading(drawing_name, level=2)

            for item in drawing_items:
                doc.add_paragraph(item.text, style="List Bullet")

        doc.save(path)
        logger.info("Word document saved: %s", path)
        return path

    # ------------------------------------------------------------------
    # PDF -- clean scope-of-work export
    # ------------------------------------------------------------------

    def generate_pdf_sync(self, path: str, **kwargs: Any) -> str:
        items: list[ClassifiedItem] = kwargs["items"]
        project_id: int = kwargs["project_id"]
        project_name: str = kwargs["project_name"]
        trade: str = kwargs["trade"]
        project_location: str = kwargs.get("project_location", "")
        drawing_s3_urls: dict[str, str] = kwargs.get("drawing_s3_urls", {})

        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        pdf_doc = SimpleDocTemplate(path, pagesize=letter)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "ScopeTitle",
            parent=styles["Title"],
            textColor=HexColor("#003366"),
            fontSize=16,
            alignment=1,  # center
        )
        info_style = ParagraphStyle(
            "ScopeInfo",
            parent=styles["Normal"],
            textColor=HexColor("#666666"),
            fontSize=11,
        )
        heading_style = ParagraphStyle(
            "ScopeHeading",
            parent=styles["Heading2"],
            textColor=HexColor("#003366"),
        )
        bullet_style = ParagraphStyle(
            "ScopeBullet",
            parent=styles["Normal"],
            leftIndent=20,
            fontSize=10,
        )

        story: list[Any] = []

        # Title
        story.append(Paragraph(f"SCOPE OF WORK \u2014 {trade.upper()}", title_style))
        story.append(Spacer(1, 12))

        # Project info
        info_parts = [f"Project: {project_name} (ID: {project_id})"]
        if project_location:
            info_parts.append(f"Location: {project_location}")
        info_parts.append(
            f"Date: {datetime.now(timezone.utc).strftime('%B %d, %Y')}",
        )
        story.append(Paragraph("<br/>".join(info_parts), info_style))
        story.append(Spacer(1, 18))

        # Scope items grouped by drawing
        grouped = _group_items_by_drawing(items)

        for drawing_name, drawing_items in grouped.items():
            s3_url = drawing_s3_urls.get(drawing_name)
            if s3_url:
                heading_text = f'<a href="{s3_url}" color="#003366">{drawing_name}</a>'
            else:
                heading_text = drawing_name
            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 6))

            for item in drawing_items:
                bullet_text = f"\u2022 {item.text}"
                story.append(Paragraph(bullet_text, bullet_style))
                story.append(Spacer(1, 3))

            story.append(Spacer(1, 6))

        pdf_doc.build(story)
        logger.info("PDF document saved: %s", path)
        return path

    # ------------------------------------------------------------------
    # CSV -- 6-column clean export
    # ------------------------------------------------------------------

    def generate_csv_sync(self, path: str, **kwargs: Any) -> str:
        items: list[ClassifiedItem] = kwargs["items"]

        header = [
            "Drawing",
            "Drawing Title",
            "Scope Item",
            "CSI Code",
            "CSI Division",
            "Trade",
            "Confidence",
        ]

        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(header)
            for item in items:
                writer.writerow([
                    item.drawing_name,
                    item.drawing_title or "",
                    item.text,
                    item.csi_code,
                    item.csi_division,
                    item.trade,
                    getattr(item, "confidence", ""),
                ])

        logger.info("CSV document saved: %s", path)
        return path

    # ------------------------------------------------------------------
    # JSON -- full data dump (UNCHANGED)
    # ------------------------------------------------------------------

    def generate_json_sync(self, path: str, **kwargs: Any) -> str:
        items: list[ClassifiedItem] = kwargs["items"]
        ambiguities: list[AmbiguityItem] = kwargs["ambiguities"]
        gotchas: list[GotchaItem] = kwargs["gotchas"]
        completeness: CompletenessReport = kwargs["completeness"]
        quality: QualityReport = kwargs["quality"]
        project_id: int = kwargs["project_id"]
        project_name: str = kwargs["project_name"]
        trade: str = kwargs["trade"]
        stats: PipelineStats = kwargs["stats"]

        output = {
            "project_id": project_id,
            "project_name": project_name,
            "trade": trade,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "items": [item.model_dump() for item in items],
            "ambiguities": [amb.model_dump() for amb in ambiguities],
            "gotchas": [gtc.model_dump() for gtc in gotchas],
            "completeness": completeness.model_dump() if completeness else None,
            "quality": quality.model_dump() if quality else None,
            "pipeline_stats": stats.model_dump() if stats else None,
        }

        with open(path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, default=str)

        logger.info("JSON document saved: %s", path)
        return path
